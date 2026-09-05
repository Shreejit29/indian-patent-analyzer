import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document


def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file.

    Args:
        file_bytes: PDF file content as bytes.

    Returns:
        Extracted text as a string.
    """

    text_parts = []

    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)

        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""

            text_parts.append(
                f"\n--- PAGE {page_number} ---\n"
                f"{page_text.strip()}\n"
            )

    except Exception as exc:
        raise ValueError(
            f"Unable to read PDF file: {exc}"
        ) from exc

    return "\n".join(text_parts).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file while preserving
    basic paragraph and table information.

    Args:
        file_bytes: DOCX file content as bytes.

    Returns:
        Extracted document text as a string.
    """

    text_parts = []

    try:
        docx_file = io.BytesIO(file_bytes)
        document = Document(docx_file)

        # Extract paragraphs
        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                text_parts.append(text)

        # Extract tables
        for table_number, table in enumerate(
            document.tables,
            start=1
        ):

            text_parts.append(
                f"\n--- TABLE {table_number} ---"
            )

            for row in table.rows:

                row_text = []

                for cell in row.cells:
                    row_text.append(
                        cell.text.strip()
                    )

                text_parts.append(
                    " | ".join(row_text)
                )

    except Exception as exc:
        raise ValueError(
            f"Unable to read DOCX file: {exc}"
        ) from exc

    return "\n".join(text_parts).strip()


def extract_text_from_file(
    file_bytes: bytes,
    filename: str
) -> str:
    """
    Automatically select the appropriate parser
    based on the uploaded file extension.

    Args:
        file_bytes: Uploaded file content.
        filename: Uploaded filename.

    Returns:
        Extracted patent text.
    """

    extension = Path(filename).suffix.lower()

    if extension == ".pdf":

        return extract_pdf_text(
            file_bytes
        )

    elif extension == ".docx":

        return extract_docx_text(
            file_bytes
        )

    else:

        raise ValueError(
            "Unsupported file format. "
            "Please upload a PDF or DOCX file."
        )


def get_document_statistics(text: str) -> dict:
    """
    Calculate basic statistics from extracted text.
    """

    words = text.split()

    paragraphs = [
        paragraph.strip()
        for paragraph in text.split("\n")
        if paragraph.strip()
    ]

    return {
        "characters": len(text),
        "words": len(words),
        "paragraphs": len(paragraphs),
    }
